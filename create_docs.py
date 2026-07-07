from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Sahifa o'lchamlari ──────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Rang konstantalari ──────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x08, 0x1a, 0x45)
ACCENT     = RGBColor(0x2d, 0x78, 0xff)
LIGHT_BLUE = RGBColor(0x57, 0xa7, 0xff)
GRAY_TEXT  = RGBColor(0x44, 0x4e, 0x6a)
WHITE      = RGBColor(0xff, 0xff, 0xff)
TABLE_HEAD = RGBColor(0x1a, 0x3a, 0x7a)
TABLE_EVEN = RGBColor(0xf0, 0xf5, 0xff)

# ── Yordamchi funksiyalar ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','bottom','left','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '2d78ff'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size  = Pt(20)
        run.font.bold  = True
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # pastki chiziq
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        btm  = OxmlElement('w:bottom')
        btm.set(qn('w:val'),   'single')
        btm.set(qn('w:sz'),    '6')
        btm.set(qn('w:space'), '4')
        btm.set(qn('w:color'), '2d78ff')
        pBdr.append(btm)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
    else:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = LIGHT_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_TEXT
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.font.bold  = True
        r1.font.color.rgb = DARK_BLUE
        r1.font.size  = Pt(11)
        r2 = p.add_run(text[len(bold_part):])
        r2.font.size  = Pt(11)
        r2.font.color.rgb = GRAY_TEXT
    else:
        r = p.add_run(text)
        r.font.size  = Pt(11)
        r.font.color.rgb = GRAY_TEXT

def info_table(rows, col_widths=None):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths is None:
        col_widths = [Cm(5.5), Cm(11)]
    for i, (k, v) in enumerate(rows):
        row = t.add_row()
        row.height = Cm(0.75)
        # kalit
        c0 = row.cells[0]
        c0.width = col_widths[0]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.bold  = True
        r0.font.size  = Pt(10)
        r0.font.color.rgb = WHITE
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(c0, '1a3a7a')
        # qiymat
        c1 = row.cells[1]
        c1.width = col_widths[1]
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size  = Pt(10)
        r1.font.color.rgb = GRAY_TEXT
        if i % 2 == 1:
            set_cell_bg(c1, 'f0f5ff')
        else:
            set_cell_bg(c1, 'ffffff')
    doc.add_paragraph()
    return t

def full_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Sarlavha qatori
    hrow = t.rows[0]
    for idx, h in enumerate(headers):
        cell = hrow.cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_bg(cell, '1a3a7a')
    # Ma'lumot qatorlari
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = 'f0f5ff' if ri % 2 == 0 else 'ffffff'
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size  = Pt(10)
            r.font.color.rgb = GRAY_TEXT
            set_cell_bg(c, bg)
    doc.add_paragraph()
    return t

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x7a)
    # fon
    pPr = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'e8f0fe')
    pPr.append(shd)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '4')
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), 'c8d8ff')
    pBdr.append(btm)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)

# ════════════════════════════════════════════════════════════════════════════
#  MUQOVA SAHIFASI
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph('\n\n')

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('Samarqand shahrining atmosfera havosi\nifloslanish geoportali')
tr.font.size  = Pt(24)
tr.font.bold  = True
tr.font.color.rgb = DARK_BLUE

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('Texnik hujjatlar / Foydalanuvchi qo\'llanmasi')
sr.font.size  = Pt(13)
sr.font.color.rgb = ACCENT

doc.add_paragraph('\n')

meta_rows = [
    ('Versiya',        'v1.0'),
    ('Sana',           '2026-yil, iyun'),
    ('Tashkilot',      'RemoteSensingInnovators'),
    ('Repository',     'github.com/RemoteSensingInnovators/dashboard_gidrometrologiya'),
    ('Ishga tushirish','python -m http.server 8000  →  localhost:8000/login.html'),
]
info_table(meta_rows)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. LOYIHA HAQIDA
# ════════════════════════════════════════════════════════════════════════════
heading('1. Loyiha haqida umumiy ma\'lumot', 1)
body(
    'Samarqand shahrining atmosfera havosi ifloslanish geoportali — bu 2016–2025 yillar '
    'davomidagi havo sifati ma\'lumotlarini interaktiv tarzda tahlil qilish va vizualizatsiya '
    'qilish uchun mo\'ljallangan veb-dastur. Tizim foydalanuvchi autentifikatsiyasini, '
    'interaktiv xaritani, uchta analitik grafikni, ma\'lumotlar jadvalini va beshta '
    'meteorologik stansiyaning KMZ-layer larini o\'z ichiga oladi.'
)

divider()

heading('1.1  Asosiy imkoniyatlar', 2)
for item in [
    ('Login tizimi — ', 'faqat ruxsat etilgan foydalanuvchilar kirishga yaxshi himoyalangan.'),
    ('Interaktiv xarita — ', 'Leaflet.js asosida Samarqand shahrining GeoJSON chegara xaritasi va 5 ta stansiya markeri.'),
    ('KMZ stansiyalar — ', 'Universitet xiyoboni, Turizm kolleji, Registon, Muzey, Dahbed stansiyalari rangli markerlar sifatida.'),
    ('Oylik trend grafigi — ', 'tanlangan yil va gaz bo\'yicha qiymatlarning oylik o\'zgarishi.'),
    ('1 yillik tahlil grafigi — ', 'barcha gazlar bo\'yicha o\'rtacha/maksimal qiymatlar bar-chart ko\'rinishida.'),
    ('Shamol yo\'nalishlari — ', 'radar-chart ko\'rinishida shamol yo\'nalishlari bo\'yicha tezliklar.'),
    ('Jadval — ', 'oylik barcha parametrlar va shamol ma\'lumotlari bitta qulay jadvalda.'),
    ('Hover tooltips — ', 'grafik ustiga sichqoncha olib borganda aniq raqamli qiymat ko\'rsatiladi.'),
]:
    bullet(item[0] + item[1], item[0])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  2. TEXNOLOGIYALAR
# ════════════════════════════════════════════════════════════════════════════
heading('2. Texnologiyalar steki', 1)

full_table(
    ['Texnologiya', 'Versiya', 'Maqsad'],
    [
        ['HTML5 / CSS3',   '—',       'Sahifa tuzilmasi va dizayn'],
        ['JavaScript ES6+','—',       'Asosiy mantiq va interaktivlik'],
        ['Leaflet.js',     '1.9.4',   'Interaktiv xarita va GeoJSON layerlar'],
        ['Chart.js',       '4.4.1',   'Line, Bar va Radar grafiklar'],
        ['Python 3.x',     '3.14+',   'Ma\'lumotlarni qayta ishlash va konvertatsiya'],
        ['python-docx',    '—',       'Hujjat yaratish'],
        ['GitHub API',     'v3',      'Fayllarni repositoryga yuklash'],
    ]
)

divider()

heading('2.1  Tashqi kutubxonalar (CDN)', 2)
full_table(
    ['Kutubxona', 'CDN manzili'],
    [
        ['Leaflet CSS', 'unpkg.com/leaflet@1.9.4/dist/leaflet.css'],
        ['Leaflet JS',  'cdn.jsdelivr.net/npm/leaflet@1.9.4/dist/leaflet.js'],
        ['Chart.js',    'cdn.jsdelivr.net/npm/chart.js@4.4.1/dist/chart.umd.min.js'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  3. FAYL TUZILMASI
# ════════════════════════════════════════════════════════════════════════════
heading('3. Fayl tuzilmasi', 1)

full_table(
    ['Fayl / Papka', 'Tavsif'],
    [
        ['index.html',              'Asosiy dashboard sahifasi'],
        ['login.html',              'Foydalanuvchi kirish sahifasi'],
        ['dashboard.js',            'Barcha grafik, xarita va jadval mantiqi'],
        ['style.css',               'Butun loyiha uchun CSS stillar'],
        ['samarqand.json',          'Samarqand shahrining GeoJSON chegarasi'],
        ['build_data.py',           'Excel fayllardan JSON ma\'lumot yaratish skripti'],
        ['convert_kmz.py',          'KMZ fayllarni GeoJSON ga o\'tkazuvchi skript'],
        ['push_to_github.py',       'GitHub API orqali fayllarni yuklash skripti'],
        ['data/samarqand_data.json','Asosiy havo sifati ma\'lumotlari (JSON)'],
        ['data/stations.geojson',   '5 ta stansiyaning GeoJSON koordinatalari'],
        ['*.kmz',                   'Google Earth KMZ formatidagi stansiya fayllari'],
        ['2016–2025.xlsx',          'Yillik havo sifati ma\'lumotlari (Excel)'],
        ['shamol.xlsx',             'Shamol yo\'nalishlari ma\'lumotlari (Excel)'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  4. LOGIN TIZIMI
# ════════════════════════════════════════════════════════════════════════════
heading('4. Login tizimi', 1)
body(
    'Tizimga kirish uchun login.html sahifasi ishlatiladi. '
    'Muvaffaqiyatli kirishdan so\'ng sessionStorage da "auth" kaliti saqlanadi. '
    'index.html har safar yuklanganida bu kalit tekshiriladi — agar yo\'q bo\'lsa, '
    'foydalanuvchi avtomatik tarzda login.html ga qaytariladi.'
)

heading('4.1  Kirish ma\'lumotlari', 2)
info_table([
    ('Foydalanuvchi nomi', 'admin'),
    ('Parol',              'admin1234'),
    ('Sessiya turi',       'sessionStorage (brauzer yopilganda o\'chadi)'),
])

heading('4.2  Autentifikatsiya oqimi', 2)
for step in [
    '1. Foydalanuvchi login.html da username va parol kiritadi.',
    '2. JavaScript tekshiradi: agar to\'g\'ri bo\'lsa, sessionStorage.setItem("auth","true") bajariladi.',
    '3. Foydalanuvchi index.html ga yo\'naltiriladi.',
    '4. index.html yuklanganda birinchi <script> bloki sessionStorage ni tekshiradi.',
    '5. "auth" kaliti yo\'q yoki "true" emas bo\'lsa — login.html ga redirect.',
    '6. "Chiqish" tugmasi: sessionStorage.removeItem("auth") va login.html ga qaytish.',
]:
    bullet(step)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  5. XARITA VA STANSIYALAR
# ════════════════════════════════════════════════════════════════════════════
heading('5. Interaktiv xarita va stansiyalar', 1)
body(
    'Xarita Leaflet.js asosida qurilgan. Sahifa yuklanganida ikkita layer avtomatik yuklanadi: '
    'Samarqand shahrining GeoJSON chegarasi va 5 ta meteorologik stansiyaning nokta layeri. '
    'Xaritaning o\'ng yuqori burchagida layer boshqaruv paneli joylashgan.'
)

heading('5.1  Stansiyalar ro\'yxati', 2)
full_table(
    ['Stansiya nomi', 'Koordinata (lon, lat)', 'Rang'],
    [
        ['Universitet xiyoboni stansiyasi', '66.9573, 39.6441', 'Qizil (#e74c3c)'],
        ['Turizm kolleji stansiyasi',       '66.9222, 39.6400', 'To\'q sariq (#e67e22)'],
        ['Registon stansiyasi',             '66.9736, 39.6530', 'Yashil (#2ecc71)'],
        ['Muzey stansiyasi',                '66.9280, 39.6666', 'Binafsha (#9b59b6)'],
        ['Dahbed stansiyasi',               '66.9725, 39.6684', 'Ko\'k (#3498db)'],
    ]
)

heading('5.2  Layer boshqaruvi', 2)
for item in [
    'Xaritadagi "Stansiyalar" checkboxni o\'chirish — barcha marker ko\'rinmaydi.',
    '"Samarqand" checkboxni o\'chirish — shahar chegarasi ko\'rinmaydi.',
    'Stansiya markeriga bosish — popup da stansiya nomi chiqadi.',
    'Stansiya markeri ustiga sichqoncha olib borish — tooltip (doimiy yozuv) ko\'rinadi.',
]:
    bullet(item)

heading('5.3  KMZ → GeoJSON konvertatsiya', 2)
body('convert_kmz.py skripti quyidagi amallarni bajaradi:')
for step in [
    'Har bir .kmz faylni ZIP arxiv sifatida ochadi.',
    'Ichidagi doc.kml faylini o\'qib XML tahlil qiladi.',
    'Placemark elementlaridan nom va koordinatalarni ajratib oladi.',
    'GeoJSON FeatureCollection formatida data/stations.geojson ga saqlaydi.',
]:
    bullet(step)

code_block('python convert_kmz.py')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  6. GRAFIKLAR
# ════════════════════════════════════════════════════════════════════════════
heading('6. Analitik grafiklar', 1)

heading('6.1  Oylik trend (Line Chart)', 2)
body('Tanlangan yil va gaz turi bo\'yicha har oylik o\'rtacha yoki maksimal qiymatni chiziq grafik sifatida ko\'rsatadi.')
info_table([
    ('X o\'qi',         'Oylar (Yanvar–Dekabr)'),
    ('Y o\'qi',         'Gaz konsentratsiyasi'),
    ('Dataset',         'Bitta gaz, bir yillik'),
    ('Hover',           'Sichqoncha ustiga borganda — oy nomi va aniq qiymat'),
])

heading('6.2  1 yillik tahlil gazlar bo\'yicha (Bar Chart)', 2)
body('Tanlangan yilda barcha gazlarning yil bo\'yi o\'rtacha qiymati ustunli diagrammada ko\'rsatiladi.')
info_table([
    ('X o\'qi',   'Gaz turlari (SO₂, NO₂, NH₃, HF, NO, Fenol, CO, Cl₂, PM)'),
    ('Y o\'qi',   'O\'rtacha konsentratsiya qiymati'),
    ('Hover',     'Sichqoncha ustiga borganda — gaz nomi va qiymat'),
])

heading('6.3  Shamol yo\'nalishlari (Radar Chart)', 2)
body('Tanlangan yil bo\'yicha har bir asosiy yo\'nalish uchun o\'rtacha shamol tezligini radar diagrammada ko\'rsatadi.')
info_table([
    ('Ko\'rinish', 'Radar (parvona)'),
    ('O\'q',       'Shamol yo\'nalishlari'),
    ('Qiymat',     'O\'rtacha tezlik (m/s)'),
    ('Hover',      'Yo\'nalish ustiga borganda — tezlik m/s da'),
])

heading('6.4  Hover tooltip dizayni', 2)
body('Barcha uchala grafik uchun bir xil tashqi HTML tooltip ishlatiladi:')
for item in [
    'Chart.js built-in tooltipini o\'chirish (enabled: false).',
    'external() callback orqali DOM da #chart-tooltip div yaratish.',
    'Sichqoncha harakati bilan tooltip pozitsiyasi yangilanadi.',
    'Oy yoki yo\'nalish nomi — kichik ko\'k sarlavha.',
    'Aniq raqamli qiymat — katta oq raqam.',
    '0.15s fade animatsiya bilan ko\'rsatiladi va yashiriladi.',
]:
    bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  7. MA'LUMOTLAR
# ════════════════════════════════════════════════════════════════════════════
heading('7. Ma\'lumotlar tuzilmasi', 1)

heading('7.1  samarqand_data.json', 2)
body('data/samarqand_data.json asosiy ma\'lumot fayli. Tuzilmasi:')
code_block(
    '{\n'
    '  "months":    ["Yanvar", "Fevral", ...],\n'
    '  "pollutants":["SO2", "NO2", "NH3", "HF", "NO", "Fenol", "CO", "CL", "Chang"],\n'
    '  "years": {\n'
    '    "2016": {\n'
    '      "Yanvar": {\n'
    '        "SO2": { "Mean": 0.012, "Max": 0.025 },\n'
    '        ...\n'
    '      }\n'
    '    }\n'
    '  },\n'
    '  "wind": {\n'
    '    "directions": ["S", "SH", "G", ...],\n'
    '    "years": { "2016": { "Yanvar": { "S": 1.2, ... } } }\n'
    '  }\n'
    '}'
)

heading('7.2  Gaz turlari', 2)
full_table(
    ['Kod', 'To\'liq nomi', 'Formula'],
    [
        ['SO2',   'Oltingugurt dioksid',  'SO₂'],
        ['NO2',   'Azot dioksid',         'NO₂'],
        ['NH3',   'Ammiak',               'NH₃'],
        ['HF',    'Vodorod ftorid',       'HF'],
        ['NO',    'Azot oksid',           'NO'],
        ['Fenol', 'Fenol',                'C₆H₅OH'],
        ['CO',    'Uglerod oksid',        'CO'],
        ['CL',    'Xlor',                 'Cl₂'],
        ['Chang', 'Chang',                'PM'],
    ]
)

heading('7.3  Ma\'lumotlarni yangilash', 2)
body('Yangi yil ma\'lumotlarini qo\'shish uchun:')
for step in [
    '1. Tegishli yil xlsx faylini loyiha papkasiga joylashtiring.',
    '2. build_data.py skriptini ishga tushiring: python build_data.py',
    '3. data/samarqand_data.json fayli avtomatik yangilanadi.',
    '4. Brauzerda sahifani yangilang (F5).',
]:
    bullet(step)
code_block('python build_data.py')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  8. ISHGA TUSHIRISH
# ════════════════════════════════════════════════════════════════════════════
heading('8. Dasturni ishga tushirish', 1)

heading('8.1  Mahalliy server (localhost)', 2)
body('Loyiha papkasiga o\'ting va Python HTTP serverini ishga tushiring:')
code_block('cd dashboard_gidrometrologiya\npython -m http.server 8000')
body('Keyin brauzerda quyidagi manzilni oching:')
code_block('http://localhost:8000/login.html')
body('Login: admin   |   Parol: admin1234', indent=True)

heading('8.2  Talablar', 2)
for item in [
    'Python 3.x (ma\'lumotlarni qayta ishlash uchun)',
    'Zamonaviy veb-brauzer (Chrome, Firefox, Edge)',
    'Internet aloqa (CDN kutubxonalar uchun: Leaflet, Chart.js)',
]:
    bullet(item)

heading('8.3  Birinchi marta ishga tushirishda', 2)
for step in [
    '1. Agar data/samarqand_data.json mavjud bo\'lmasa: python build_data.py',
    '2. Agar data/stations.geojson mavjud bo\'lmasa: python convert_kmz.py',
    '3. python -m http.server 8000',
    '4. http://localhost:8000/login.html',
]:
    bullet(step)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  9. GITHUB
# ════════════════════════════════════════════════════════════════════════════
heading('9. GitHub repository', 1)

info_table([
    ('Repository URL', 'https://github.com/RemoteSensingInnovators/dashboard_gidrometrologiya'),
    ('Branch',         'main'),
    ('Tashkilot',      'RemoteSensingInnovators'),
    ('Litsenziya',     'MIT'),
])

heading('9.1  Yangilashlar tarixi', 2)
full_table(
    ['Commit', 'O\'zgartirish'],
    [
        ['Dastlabki',  'Asosiy dashboard: xarita, grafiklar, jadval, ma\'lumotlar'],
        ['Update 1',   '5 ta KMZ stansiya layeri xaritaga qo\'shildi'],
        ['Update 2',   'Login tizimi (admin/admin1234) qo\'shildi'],
        ['Update 3',   'Chart tooltip dizayni yaxshilandi (HTML external tooltip)'],
        ['Update 4',   'Chart sarlavha: "1 yillik tahlil gazlar bo\'yicha"'],
        ['Update 5',   'Chiqish tugmasi dizayni yangilandi'],
    ]
)

heading('9.2  GitHub API orqali yuklash', 2)
body('push_to_github.py skripti Windows Credential Manager da saqlangan tokendan foydalanadi. Yangi fayllarni yuklash uchun:')
code_block('python push_to_github.py')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  10. UI / DIZAYN
# ════════════════════════════════════════════════════════════════════════════
heading('10. UI va dizayn', 1)
body('Butun interfeys quyuq ko\'k-temir rang sxemasida yaratilgan:')

full_table(
    ['Element', 'CSS qiymati', 'Tavsif'],
    [
        ['Fon',             '#091528 → #02040a',  'Vertikal gradient'],
        ['Asosiy matn',     '#eef3ff',             'Oq-ko\'k'],
        ['Accent rang',     '#57a7ff',             'Ko\'k yoritilgan'],
        ['Card fon',        'rgba(15,26,48,0.95)', 'Shisha effekti'],
        ['Stansiya label',  'rgba(8,16,40,0.88)',  'Quyuq tooltip'],
        ['Tooltip fon',     'rgba(12,22,50,0.96)', 'Chart tooltip'],
        ['Chiqish tugmasi', 'rgba(255,255,255,0.04)', 'Hover da qizil'],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
#  OXIRGI SAHIFA
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_paragraph('\n\n\n')
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = end_p.add_run('RemoteSensingInnovators\nSamarqand shahrining atmosfera havosi ifloslanish geoportali\n2026')
er.font.size  = Pt(12)
er.font.color.rgb = GRAY_TEXT

# ── Saqlash ─────────────────────────────────────────────────────────────────
out = 'Geoportal_Hujjatlar.docx'
doc.save(out)
print('OK: ' + out)
